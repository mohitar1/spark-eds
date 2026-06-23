#!/usr/bin/env python3
"""Generate DA-ready .docx files for the Frescopa-branded spark-eds asset portal.

Landing + dedicated search pages. Pages map to DA paths under en/:
  index.docx              -> en/index             (landing: hero search-bar + Browse-by-category)
  search.docx             -> en/search            (search-bar + asset results)
  search-collections.docx -> en/search-collections (search-bar + collection results)
  collection-details.docx -> en/collection-details (single collection asset results)
  about.docx              -> en/about             (brand story)
  nav.docx                -> en/nav               (header fragment: brand + tools)
  footer.docx             -> en/footer            (footer fragment)
  logins.docx        -> en/reports/logins         (admin users report)
  report-hub.docx    -> en/reports/report-hub      (reports landing page)
  searches.docx      -> en/reports/searches
  asset-activity.docx -> en/reports/asset-activity (asset use audit)

IMPORTANT (DA workflow):
  Upload each .docx DIRECTLY into DA (drag into the DA file browser). Do NOT open
  or save these files in Microsoft Word first -- Word rewrites the relative links
  as file:// paths and collapses the section breaks, which breaks the pages.
"""

import json
from pathlib import Path
from urllib.parse import quote

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

OUT = Path(__file__).parent

LOGO = ':frescopa-icon:'

ASSET_EXC_FACETS = """{
"dc:subject": {
"label": "Keywords",
"type": "string"
},
"dc:format": {
"label": "Format",
"type": "string"
}
}"""

COLLECTION_EXC_FACETS = """{
"dc:subject": {
"label": "Keywords",
"type": "string"
},
"repo:createdBy": {
"label": "Uploaded By",
"type": "string"
}
}"""

BRAND_STORY = [
    'Born from a passion for exceptional coffee and modern living, Fréscopa blends '
    'the craftsmanship of classic espresso with the convenience and flair of today\u2019s '
    'lifestyle. Every product, every moment, beautifully captured.',
    'This portal is the home of the Fr\u00e9scopa brand library \u2014 search, preview, '
    'and download product photography, machine imagery, lifestyle shots, packaging, and more.',
]

CATEGORY_FACET_KEY = 'dc:subject'

CATEGORY_TILES = [
    ('Coffee', 'Bagged coffee, pods, and signature roasts.', 'coffee'),
    ('Machines', 'Espresso machines and smart brewing gear.', 'machine'),
    ('Accessories', 'Grinders, thermoses, and add-ons.', 'accessory'),
    ('Lifestyle', 'In-the-moment brand and lifestyle imagery.', 'lifestyle'),
]


def category_search_url(facet_value):
    """Build a search URL that applies a dc:subject facet filter (not full-text query)."""
    facet_filters = json.dumps(
        {CATEGORY_FACET_KEY: {facet_value: True}},
        separators=(',', ':'),
    )
    return f'/en/search?facetFilters={quote(facet_filters)}'


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hyperlink(paragraph, text, url, bold=False, italic=False):
    """Insert a real external hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_button(doc, text, url, secondary=False):
    """A standalone single-link paragraph becomes an EDS button."""
    p = doc.add_paragraph()
    add_hyperlink(p, text, url, bold=not secondary, italic=secondary)
    return p


def add_block_table(doc, block_name, rows=None, cols=2):
    rows = rows if rows is not None else [['', '']]
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = block_name
    for i, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            table.rows[i].cells[c].text = value
    doc.add_paragraph('')


def add_metadata_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Metadata'
    for i, (a, b) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b


def add_section_metadata(doc, style):
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Section Metadata'
    table.rows[1].cells[0].text = 'style'
    table.rows[1].cells[1].text = style
    doc.add_paragraph('')


def add_nav_role_metadata(doc, role):
    """Mark a nav fragment section as brand, sections, or tools (independent slots)."""
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Section Metadata'
    table.rows[1].cells[0].text = 'role'
    table.rows[1].cells[1].text = role
    doc.add_paragraph('')
def add_cards(doc, cards):
    """1-column Cards block: each row is a category tile with a Browse link."""
    table = doc.add_table(rows=1 + len(cards), cols=1)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'cards'
    for i, (title, blurb, facet_value) in enumerate(cards, start=1):
        cell = table.rows[i].cells[0]
        cell.paragraphs[0].text = ''
        head = cell.paragraphs[0]
        head.style = doc.styles['Heading 3']
        head.add_run(title)
        cell.add_paragraph(blurb)
        link_p = cell.add_paragraph()
        add_hyperlink(link_p, 'Browse \u2192', category_search_url(facet_value), bold=True)
    doc.add_paragraph('')


def build_index():
    """Landing: branded hero (search bar) + Browse-by-category."""
    doc = Document()

    # Section 1: compact branded hero + search bar (no results grid)
    doc.add_paragraph(LOGO)
    doc.add_heading('Find your Fr\u00e9scopa assets', level=1)
    doc.add_paragraph(
        'Search, preview, and download brand-approved coffee, machine, and lifestyle imagery.'
    )
    add_block_table(doc, 'search-bar')
    add_section_metadata(doc, 'search-hero')
    add_hr(doc)

    # Section 2: Browse by category tiles
    doc.add_heading('Browse by category', level=2)
    add_cards(doc, CATEGORY_TILES)
    add_section_metadata(doc, 'category-tiles')
    add_hr(doc)

    add_metadata_table(doc, [
        ('title', 'Fr\u00e9scopa Asset Library'),
        ('description', 'Search, preview, and download brand-approved Fr\u00e9scopa assets.'),
    ])

    path = OUT / 'index.docx'
    doc.save(path)
    print(f'Wrote {path}  -> upload into DA: mohitar1/spark-eds -> en/index')


def build_search():
    """Dedicated asset search page: search bar + results grid (Santander pattern)."""
    doc = Document()

    add_block_table(doc, 'search-bar')
    add_hr(doc)
    add_block_table(doc, 'search-results', [
        ('excFacets', ASSET_EXC_FACETS),
    ])
    add_hr(doc)
    add_metadata_table(doc, [
        ('title', 'Search Assets'),
        ('description', 'Search, browse, and download Fr\u00e9scopa brand assets.'),
    ])

    path = OUT / 'search.docx'
    doc.save(path)
    print(f'Wrote {path}  -> upload into DA: mohitar1/spark-eds -> en/search')


def build_search_collections():
    """Dedicated collection search page: search bar + collection results (Santander pattern)."""
    doc = Document()

    add_block_table(doc, 'search-bar')
    add_hr(doc)
    add_block_table(doc, 'search-collection-results', [
        ('excFacets', COLLECTION_EXC_FACETS),
    ])
    add_hr(doc)
    add_metadata_table(doc, [
        ('title', 'Search Collections'),
        ('description', 'Search and browse Fr\u00e9scopa asset collections.'),
    ])

    path = OUT / 'search-collections.docx'
    doc.save(path)
    print(f'Wrote {path}  -> upload into DA: mohitar1/spark-eds -> en/search-collections')


def build_collection_details():
    """Single collection detail page: collection-details block reads the collection id from the URL."""
    doc = Document()

    add_block_table(doc, 'collection-details', [
        ('excFacets', COLLECTION_EXC_FACETS),
    ])
    add_hr(doc)
    add_metadata_table(doc, [
        ('title', 'Collection Details'),
        ('description', 'View and browse assets in a Fréscopa collection.'),
    ])

    path = OUT / 'collection-details.docx'
    doc.save(path)
    print(f'Wrote {path}  -> upload into DA: mohitar1/spark-eds -> en/collection-details')


def build_about():
    doc = Document()

    doc.add_heading('About Fr\u00e9scopa', level=1)
    for para in BRAND_STORY:
        doc.add_paragraph(para)
    add_button(doc, 'Browse the library', '/en/search')
    add_hr(doc)
    add_metadata_table(doc, [
        ('title', 'About Fr\u00e9scopa'),
        ('description', 'The story behind the Fr\u00e9scopa brand and asset library.'),
    ])

    path = OUT / 'about.docx'
    doc.save(path)
    print(f'Wrote {path}  -> upload into DA: mohitar1/spark-eds -> en/about')


def build_nav():
    doc = Document()

    # Two independent sections (HR-separated): brand (logo) and tools (icon slot).
    # Cart / download / bell icons are injected into tools by header.js.

    # Section 1 — brand
    brand = doc.add_paragraph()
    add_hyperlink(brand, LOGO, '/en/')
    add_nav_role_metadata(doc, 'brand')
    add_hr(doc)

    # Section 2 — tools (do not merge with brand)
    doc.add_paragraph('\u00a0')
    add_nav_role_metadata(doc, 'tools')

    path = OUT / 'nav.docx'
    doc.save(path)
    print(f'Wrote {path}  -> upload into DA: mohitar1/spark-eds -> en/nav')


def build_report_page(filename, block_name, title, description, da_path):
    """Single-section admin report page with one block + page metadata."""
    doc = Document()

    add_block_table(doc, block_name)
    add_hr(doc)
    add_metadata_table(doc, [
        ('title', title),
        ('description', description),
    ])

    path = OUT / filename
    doc.save(path)
    print(f'Wrote {path}  -> upload into DA: mohitar1/spark-eds -> {da_path}')


def build_reports_hub():
    build_report_page(
        'report-hub.docx',
        'report-hub',
        'Reports',
        'System reports hub for administrators.',
        'en/reports/report-hub',
    )


def build_reports_logins():
    build_report_page(
        'logins.docx',
        'report-logins',
        'Users Report',
        'User activity and engagement report for administrators.',
        'en/reports/logins',
    )


def build_reports_searches():
    build_report_page(
        'searches.docx',
        'report-searches',
        'Search Analytics',
        'Search activity report for administrators.',
        'en/reports/searches',
    )


def build_reports_asset_activity():
    build_report_page(
        'asset-activity.docx',
        'report-asset-activity',
        'Asset Activity',
        'Asset views, downloads, and usage audit for administrators.',
        'en/reports/asset-activity',
    )


def build_my_notifications():
    """User notifications inbox page (my-notifications block)."""
    build_report_page(
        'my-notifications.docx',
        'my-notifications',
        'My Notifications',
        'Your notifications inbox — system notices and personal alerts.',
        'en/my-dam/my-notifications',
    )


def build_footer():
    doc = Document()

    # Section 1: brand column (logo stacked over tagline)
    doc.add_paragraph(LOGO)
    doc.add_paragraph('Premium coffee, beautifully captured.')
    doc.add_paragraph('Your brand asset distribution portal.')
    add_hr(doc)

    # Section 2: link columns
    columns = [
        ('Assets', [('Search Assets', '/en/search'), ('Browse Collections', '/en/search-collections')]),
        ('Company', [('About', '/en/about'), ('Contact', 'mailto:assets@frescopa.coffee')]),
        ('Help', [('Documentation', '/en/about'), ('Support', 'mailto:assets@frescopa.coffee')]),
    ]
    table = doc.add_table(rows=2, cols=len(columns))
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'columns'
    for i, (header, links) in enumerate(columns):
        cell = table.rows[1].cells[i]
        cell.paragraphs[0].text = ''
        head = cell.paragraphs[0]
        head.add_run(header).bold = True
        for label, url in links:
            link_p = cell.add_paragraph()
            add_hyperlink(link_p, label, url)
    doc.add_paragraph('')
    add_hr(doc)

    # Section 3: copyright
    doc.add_paragraph('\u00a9 2026 Fr\u00e9scopa. All rights reserved.')

    path = OUT / 'footer.docx'
    doc.save(path)
    print(f'Wrote {path}  -> upload into DA: mohitar1/spark-eds -> en/footer')


def main():
    build_index()
    build_search()
    build_search_collections()
    build_collection_details()
    build_about()
    build_reports_hub()
    build_reports_logins()
    build_reports_searches()
    build_reports_asset_activity()
    build_my_notifications()
    build_nav()
    build_footer()
    print(
        '\nDone. Upload each .docx DIRECTLY into https://da.live/edit#/mohitar1/spark-eds'
        '\n(do NOT open/save them in Microsoft Word first).'
        '\n\nReports: create folder en/reports/ in DA, then upload reports-*.docx pages there.'
    )


if __name__ == '__main__':
    main()
